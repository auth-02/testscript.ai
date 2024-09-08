from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from typing import List
import base64
import google.generativeai as genai
import os
from PIL import Image
from dotenv import load_dotenv
from io import BytesIO
import json
from docx import Document
from io import BytesIO
from PIL import Image as PilImage
from docx.shared import Inches
from docx2pdf import convert

load_dotenv()
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def create_test_case_doc(test_cases, img_dir='D:\\CODESPACES\\testscript.ai\\tests\\img'):
    doc = Document()

    print(f"Type of test_cases: {type(test_cases)}")
    print(f"Content of test_cases: {test_cases}")
    
    cases = test_cases["test_cases"]
    if not isinstance(cases, list):
        raise TypeError("Expected 'test_cases' to be a list")

    for i, case in enumerate(cases, 1):
        if not isinstance(case, dict):
            raise TypeError(f"Expected case to be a dictionary, got {type(case)}")
        
        doc.add_heading(f"Test Case ID: TC{i:03d}", level=1)
        doc.add_paragraph(f"Test Case Title: {case.get('test_case_title', 'N/A')}")
        
        doc.add_heading('Description:', level=2)
        doc.add_paragraph(case.get('description', 'N/A'))
        
        doc.add_heading('Preconditions:', level=2)
        doc.add_paragraph(case.get('pre_conditions', 'N/A'))
        
        doc.add_heading('Test Steps:', level=2)
        steps = case.get('steps', [])
        if steps:
            for j, step in enumerate(steps, 1):
                doc.add_paragraph(f"{j}. {step}")
        else:
            doc.add_paragraph("N/A")
        
        doc.add_heading('Expected Results:', level=2)
        doc.add_paragraph(case.get('expected_result', 'N/A'))
        
        doc.add_heading('Actual Results:', level=2)
        doc.add_paragraph(case.get('actual_result', 'Not Executed'))
        
        doc.add_heading('Test Data:', level=2)
        test_data = case.get('test_data', {})
        if isinstance(test_data, dict) and test_data:
            for key, value in test_data.items():
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value if value else 'N/A'}")
        else:
            doc.add_paragraph("N/A")
        
        doc.add_heading('Notes:', level=2)
        doc.add_paragraph(case.get('notes', 'N/A'))

        # img_path = os.path.join(img_dir, f"{i}_unnamed.png")
        # if os.path.exists(img_path):
        #     with PilImage.open(img_path) as img:
        #         img.verify() 
        #         img = img.resize((300, 200)) 
        #         img_buffer = BytesIO()
        #         img.save(img_buffer, format='PNG')
        #         img_buffer.seek(0)
        #         doc.add_paragraph(f"Screenshot {i}")
        #         doc.add_picture(img_buffer, width=Inches(3)) 
        
        doc.add_page_break()

    doc_filename = "test_cases.docx"
    doc.save(doc_filename)
    
    return doc_filename

def convert_doc_to_pdf(doc_filename):
    pdf_filename = os.path.splitext(doc_filename)[0] + '.pdf'
    convert(doc_filename, pdf_filename)
    return pdf_filename

@app.post("/generate-test-cases")
async def generate_test_cases(context: str = Form(...), images: List[UploadFile] = File(...)):
    
    print("")
    print(f"Received context: {context}")
    print("")
    print("Received Images:")
    for image in images:
        print(f"Image filename: {image.filename}, content type: {image.content_type}")
        
    print(f"Received {len(images)} images in total")
    
    image_parts = []
    for image in images:
        content = await image.read()
        image_parts.append({
            "mime_type": image.content_type,
            "data": base64.b64encode(content).decode('utf-8')
        })

    prompt = """Based on the following screenshots and context, generate detailed test cases in the structured JSON format below for each functionality shown. 
The JSON format should include the following fields for each test case:

- "test_case_title": A brief title summarizing the functionality being tested
- "description": A brief description explaining what the test case is verifying
- "pre_conditions": The setup or conditions that must be met before running the test
- "steps": An array of step-by-step instructions on how to carry out the test
- "expected_result": The expected outcome if the functionality is working correctly
- "actual_result": The actual observed outcome of the test (set as what you think might be the most correct for the test case)
- "test_data": A JSON object containing any specific input data required to perform the test
- "notes": Additional information or edge cases to consider (set as what you think might be the most correct for the test case)

If any field is not applicable, please set its value to null.

The JSON structure should look like this:

{
    "test_cases": [
        {
            "test_case_title": "Verify Login Functionality",
            "description": "This test case will verify the login functionality of the application.",
            "pre_conditions": "User must have a valid account",
            "steps": [
                "Open the app",
                "Enter username",
                "Enter password",
                "Click the 'Login' button"
            ],
            "expected_result": "User is successfully logged in",
            "actual_result": "Upon entering the username & password, user must be able to login into the app successfully",
            "test_data": {
                "username": "testuser",
                "password": "testpassword"
            },
            "notes": "Ensure the password is case-sensitive and handled correctly. Test the behavior with invalid login attempts and edge cases like empty fields."
        },
        {
            "test_case_title": "Verify Signup Functionality",
            "description": "This test case will verify the signup functionality of the application.",
            "pre_conditions": "User must not have an account",
            "steps": [
                "Open the app",
                "Enter username",
                "Enter email",
                "Enter password",
                "Click the 'Signup' button"
            ],
            "expected_result": "User is successfully signed up",
            "actual_result": "Upon entering the username & password, user must be able to signup into the app successfully",
            "test_data": {
                "username": "newuser",
                "email": "newuser@example.com",
                "password": "newpassword"
            },
            "notes": "Ensure the email follows the correct format and that the password meets any complexity requirements. Test for duplicate usernames or emails and check for appropriate error handling."
        }
    ]
}


Context: {context}
If any field is not applicable, please set its value to null.
Please provide step-by-step instructions for testing each feature visible in the screenshots in the JSON format, ensuring that any missing or inapplicable fields are set to None. Do not include ```json at the start and ``` at the end of the response."""


    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([prompt, *image_parts])
    
    try:
        print("response.text", response.text)
        # test_cases = response.text
        test_cases = json.loads(response.text)
        print("test_cases", test_cases)
    except Exception as e:
        test_cases = [{"description": "Error in generating structured test cases", "pre_conditions": "N/A", "steps": ["Review the generated content manually"], "expected_result": "N/A"}]
        print(e)

    doc_filename = create_test_case_doc(test_cases, images)
    pdf_filename = convert_doc_to_pdf(doc_filename)
    
    return FileResponse(pdf_filename, media_type="application/pdf", filename="test_instructions.pdf")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
